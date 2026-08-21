#!/usr/bin/env python3
"""Trim Chapters 3 and 6 of Draft 7. Duplication and detail out, claims kept."""
import os
import docx

P = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/UCL Compling/Dissertation/Aug 6/Dissert Draft 7.docx")

# (opening words used to locate the paragraph, replacement text)
TRIMS = [
# ---------------- 3.2 ----------------
("The material is naturalistic speech from political podcasts",
 "The material is naturalistic speech from political podcasts and broadcast programmes, chosen over "
 "acted emotion corpora because acted and spontaneous renditions of the same category are not "
 "interchangeable stimuli. Scherer (2003) sets that problem out for vocal emotion research generally "
 "and Rockwell (2000) demonstrates it for sarcasm, finding listeners able to discriminate posed "
 "sarcasm but not spontaneous sarcasm. Naturalistic podcast-derived corpora were built in response "
 "(Busso et al., 2025), and this corpus follows that precedent at a finer grain, since the unit of "
 "analysis is a phrase rather than a speaking turn."),

("The study targets eight short high-frequency phrases",
 "The study targets eight short high-frequency phrases whose pragmatic force varies by delivery while "
 "their lexical content is fixed, namely yeah, okay, right, sure, great, fine, really and come on. The "
 "four agreement particles are response tokens in the conversation-analytic sense, a class Gardner "
 "(2001) characterises as carrying no dictionary meaning but conveying the listener's stance through "
 "phonetic form, prosodic shape and placement. Gravano et al. (2012) add that their functions are "
 "separated in spontaneous dialogue by prosodic realisation rather than by wording, which predicts the "
 "distribution in [B.5] where neutral stance concentrates almost entirely in these four words. The "
 "evaluative terms and challenge markers are grouped on functional grounds for this study."),

("Three nested context windows were cut",
 "Three nested context windows were cut for each occurrence, centred on the target word midpoint, at "
 "6, 10 and 16 s, corresponding approximately to the target segment alone, the segment with one "
 "neighbour either side, and a full conversational turn either side. The primary window throughout is "
 "the 10 s segment window. Cutting parameters and normalisation are in Appendix [B]."),

# ---------------- 3.5 ----------------
("Why these. Mimi is the object of study",
 "Why these. Mimi is the object of study because it is the tokeniser a deployed full-duplex system "
 "consumes rather than a research codec [2.2], and it distils its first codebook from WavLM-Large "
 "(Défossez et al., 2024), so WavLM is not a free choice. The Whisper encoder tests whether a "
 "transcription objective strips paralinguistic content. Its small variant is the one genuinely "
 "constrained choice, and the constraint bears asymmetrically, since a null result would have been "
 "uninterpretable while the positive result in [4.2] can only have been worked against by limited "
 "capacity."),

("Three codecs are compared rather than two",
 "Three codecs are compared rather than two, and the third is what makes the architectural comparison "
 "possible. DAC (Kumar et al., 2023) is the undistilled contrast, trained on reconstruction and "
 "adversarial objectives alone. EnCodec (Défossez et al., 2022) matches DAC's 75 Hz frame rate while "
 "differing in encoder architecture, which turns a two-point comparison into a three-point one with "
 "frame rate held constant. MPNet supplies the text baseline, chosen because the argument requires a "
 "strong text competitor rather than a strawman. eGeMAPS supplies a hand-crafted acoustic baseline "
 "whose features are individually interpretable and can therefore be grouped by what they measure, "
 "which the cue analysis in [4.4] requires. HuBERT-large is reported in Appendix [H]."),

("That comparison requires care about which vectors are commensurate",
 "That comparison requires care about which vectors are commensurate. In all three codecs the residual "
 "quantiser operates in a projected space, so comparing the encoder latent against the quantiser's "
 "reconstructed output would compare vectors that are not aligned, their cosine over all 873 clips "
 "being 0.004. The comparable pair is the projected latent against the summed codebook vectors taken "
 "before output projection, whose cosine is 0.821 for Mimi. EnCodec needs no such correction. "
 "Per-codec figures are in Appendix [B]."),

# ---------------- 3.7 ----------------
("Every reported figure is the mean over 25 independent",
 "Every reported figure is the mean over 25 independent episode-to-fold partitions, constructed within "
 "the analysis code rather than delegated to a library. Two considerations require this. Bengio and "
 "Grandvalet (2004) show there is no universal unbiased estimator of the variance of K-fold "
 "cross-validation, so the uncertainty of a single run cannot be recovered from that run. Bouthillier "
 "et al. (2021) recommend randomising as many sources of variation as possible, since averaging over "
 "an imperfect estimator approaches the ideal one far more cheaply than reducing any single source. "
 "Measured here, fold assignment alone contributes a standard deviation of 0.010 with a range up to "
 "0.06, the magnitude of several differences this study reports, and the assignment produced by a "
 "standard grouped splitter changed between two library releases, moving one representation by 0.020 "
 "on identical inputs."),

("Layer selection. For the continuous encoders",
 "Layer selection. For the continuous encoders a single layer is carried into the main analyses, fixed "
 "in advance of the timing measurements by taking the maximum of the pooled curve reported by the "
 "initial probe battery, giving L20 for WavLM, L9 for Whisper and L23 for HuBERT. Two properties of "
 "that rule are measured rather than assumed. Choosing a layer by maximising over the same data that "
 "reports its score is optimistic, and the choice is not stable across folds for HuBERT, which is why "
 "it is reported separately [H, G.6]. The layers were fixed before the readout comparisons existed, so "
 "they are not selected to favour any result in [4.4], and the full 63-layer sweep is in [G.6]."),

("Probe capacity. A linear probe measures",
 "Probe capacity. A linear probe measures whether information is linearly accessible, which is not the "
 "same as whether it is present (Belinkov, 2022), and a probe may succeed by memorising rather than "
 "reading structure, which Hewitt and Liang (2019) address through control tasks with randomised "
 "labels, of which the permutation null is an instance. Accessibility is tested rather than argued "
 "around, by a strongly regularised non-linear probe under six capacity settings, with the bound "
 "across that sweep carried forward [G.3]."),

("The analyses. Chapter 4 reports five.",
 "The analyses. Chapter 4 reports five: pooled decodability with the within-word contrast that is the "
 "lexical control [4.2], the quantisation ladder across three codecs [4.3], cue retention and the "
 "order effect against the shuffled control [4.4], the architectural comparison at matched frame rate "
 "[4.5], and the controls [4.6]."),

# ---------------- 6.1 ----------------
("This study asked whether pragmatic contrast survives",
 "This study asked whether pragmatic contrast survives the representations that speech-to-speech "
 "systems consume, and if not, where it is lost. Of the five hypotheses formed before the analyses, "
 "H1, H2, H3 and H5 are supported and H4 is not. Of the four formed after the falsification of H4 "
 "redirected the study, H9 is supported and H6 to H8 are not. As [1.3] states, H9 was formed after an "
 "anomalous result and then tested on a codec not previously included, which is not the same kind of "
 "claim as one registered in advance."),

("Stance is linearly decodable under lexical control",
 "Stance is linearly decodable under lexical control. WavLM reaches 0.557 macro-F1 against its own "
 "null of 0.330, and the ordering survives the within-word contrast, matched arousal, held-out shows "
 "and six probe capacities [4.2, 4.6]. This is decodability relative to other representations rather "
 "than to a listener, since the best model condition reaches 0.533 where annotators reach 0.730 [4.1]."),

("What the encoder fails to preserve is temporal organisation",
 "What the encoder fails to preserve is temporal organisation rather than acoustic detail. The codecs "
 "recover hand-crafted cues more faithfully than the model Mimi distils from, including the contour "
 "features stance is built from, while reading stance off them far worse [4.4]. Temporal is the one "
 "cue group they lose, and independently the gain from frame order declines across the encoder rungs. "
 "Matching DAC and EnCodec at 75 Hz isolates the reason [4.5]."),

("A fifth result bears on the literature",
 "A fifth result bears on the literature rather than on codecs. Fold assignment alone contributes a "
 "standard deviation of 0.010 with a range up to 0.06, the magnitude of several differences reported "
 "in this area [3.7], so differences under roughly 0.03 are not read as orderings here."),

# ---------------- 6.2 ----------------
("Three hypotheses were stated before testing and are not supported",
 "Three hypotheses were stated before testing and are not supported, and reporting them is part of the "
 "result. Variable-frame-rate tokenisation does not recover the loss. Timing features alone sit at "
 "chance. Order-aware summaries of the discrete streams score below the unigram histogram they were "
 "intended to improve on [4.6]. A training-free contrast-preservation measure was also attempted and "
 "does not discriminate, and its failure is structural rather than a matter of power [G.5]."),

("The same caveat bounds the architectural account",
 "The architectural account cannot be separated from training history. Frozen public checkpoints "
 "cannot distinguish what an architecture can represent from what a particular training run taught it "
 "to represent, and an apparent architectural limit in DAC has already been shown to resolve into a "
 "training misconfiguration (Gichamba and Busogi, 2026)."),

("The encoder figure is an upper bound",
 "The encoder figure is an upper bound while the quantisation figure is not, since WavLM contributes "
 "twice the pooled dimensionality of a codec latent while the quantiser step compares a representation "
 "against itself at equal width. That matters because the ratio between the two is the claim [F]."),

# ---------------- 6.3 ----------------
("Test the organisation account directly",
 "Test the organisation account directly. Which cues a codec loses is now measured and the answer is "
 "essentially none of them, so the open question is why surviving acoustics are not usable. The "
 "sharpest test holds a representation fixed in acoustic content and varies only how that content is "
 "arranged, for instance by reorganising a codec latent under a contrastive objective with no access "
 "to new signal. Same-word opposite-stance pairs make delivery the only separating signal, and the 873 "
 "clips assembled here are already in that form."),

("Build the diagnostic that would notice",
 "Build the diagnostic that would notice. The retention that exists is optimised for by no loss term, "
 "and reconstruction quality, word error rate and perceptual scores would be unchanged if it vanished, "
 "so a successor system could drop the objective and every reported number would improve. The "
 "minimal-pair ABX task is the natural template, since it fixes the comparison as a triplet and "
 "estimates no centroid, so it degrades gracefully at small samples (Schatz et al., 2013) and is used "
 "this way inside evaluation suites for spoken language modelling (Dunbar et al., 2021). A "
 "pragmatic-contrast analogue would present two clips of one word carrying opposing stance and a third "
 "matched to one of them. Preservation that appears on a scorecard is preservation that can be "
 "defended."),
]


def main():
    d = docx.Document(P)
    before = sum(len(p.text.split()) for p in d.paragraphs)
    print("Chapters 3 and 6\n" + "-" * 62)
    saved = 0
    for anchor, new in TRIMS:
        for p in d.paragraphs:
            t = p.text.replace("\xa0", " ").strip()
            if t.startswith(anchor):
                old_n = len(t.split()); new_n = len(new.split())
                if not p.runs: p.add_run("")
                p.runs[0].text = new
                for r in p.runs[1:]: r.text = ""
                saved += old_n - new_n
                print(f"  {old_n:>4} -> {new_n:<4} ({old_n-new_n:+4})  {anchor[:44]}")
                break
        else:
            print(f"  ** NOT FOUND **  {anchor[:50]}")
    d.save(P)
    after = sum(len(p.text.split()) for p in docx.Document(P).paragraphs)
    print("-" * 62)
    print(f"  saved {saved} words   document {before} -> {after}")


if __name__ == "__main__":
    main()
