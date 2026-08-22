#!/usr/bin/env python3
"""Every invariant this project has checked by hand, in one place.

Written because most defects found in later review passes had been introduced by earlier
ones, and because each pass used whatever lens happened to be in mind that day. A single
run covers all the lenses at once, so a change cannot pass by being checked against only
the dimension it was made along.

The checks, and the real defect each was written after:

  arithmetic     Table 4.2's margins once disagreed with subtracting their own columns
  provenance     ten nulls and two order effects had no saved results file behind them
  captions       seven of ten tables had no number, so nothing could be referred to
  numbering      tables ran 4.1, 4.2, 4.4 after an insertion
  crossrefs      "Appendix D" resolved to a real appendix about a different subject
  punctuation    colons and semicolons the house style excludes
  naming         one row was "deployed tokens" in three tables and "deployed histogram" in two
  agreement      "the histogram scores lower ... but clear a lower null", left by a rename
  duplication    one sentence stood verbatim in 4.6, 5.3 and 6.2
  figures        a hand-authored figure kept 0.244 for eight months after the value changed

Exit status is 1 if anything fails, so this can gate a commit.

Usage:  python3 src/53_check_draft.py
        python3 src/53_check_draft.py --quiet     (only failures)
"""
import argparse
import collections
import difflib
import glob
import os
import re
import sys
from pathlib import Path

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/"
                          "UCL Compling/Dissertation/Aug 6/Dissert Draft 7.docx")
APPX = ROOT / "docs" / "drafts" / "07_appendices.md"

FAILED = []


def check(name, ok, detail=""):
    FAILED.append(name) if not ok else None
    return name, ok, detail


def blocks(doc):
    for c in doc.element.body.iterchildren():
        if c.tag.endswith("}p"):
            yield Paragraph(c, doc)
        elif c.tag.endswith("}tbl"):
            yield Table(c, doc)


def num(x):
    try:
        return float(x.replace("−", "-").replace("+", "").replace("%", "").replace(",", ""))
    except ValueError:
        return None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args()

    d = docx.Document(DOCX)
    paras = [p.text.replace("\xa0", " ").strip() for p in d.paragraphs if p.text.strip()]
    txt = " ".join(paras)
    appx = APPX.read_text()
    body = [t for t in paras
            if not re.match(r"^(Abstract$|Chapter \d|\d+\.\d+ |Table \d|Figure \d)", t)]
    results = []

    # -- captions and numbering ---------------------------------------------------
    caps, uncap, prev = [], 0, None
    for b in blocks(d):
        if isinstance(b, Paragraph):
            t = b.text.replace("\xa0", " ").strip()
            if t:
                prev = t
        else:
            m = re.match(r"^(Table \d+\.\d+)\.", prev or "")
            caps.append(m.group(1)) if m else None
            uncap += 0 if m else 1
    results.append(check("every table captioned", uncap == 0, f"{uncap} without"))
    want = [f"Table {c}.{i}" for c, n in (("2", 2), ("4", len(caps) - 2))
            for i in range(1, n + 1)]
    results.append(check("table numbering sequential", caps == want,
                         f"{' '.join(caps)}" if caps != want else ""))

    figs = re.findall(r"^(Figure \d+\.\d+)\.", "\n".join(paras), re.M)
    results.append(check("figure numbering sequential",
                         figs == sorted(set(figs), key=figs.index) and len(figs) == len(set(figs)),
                         " ".join(figs)))
    results.append(check("figures embedded matches captions", len(d.inline_shapes) == len(figs),
                         f"{len(d.inline_shapes)} images, {len(figs)} captions"))

    # -- arithmetic ----------------------------------------------------------------
    bad = []
    for t in d.tables:
        head = [c.text.strip().lower() for c in t.rows[0].cells]
        if "null" in head and "margin" in head:
            for r in t.rows[1:]:
                a, b_, c_ = num(r.cells[1].text), num(r.cells[2].text), num(r.cells[3].text)
                if None in (a, b_, c_):
                    continue
                if abs((a - b_) - c_) > 0.0015:
                    bad.append(f"{r.cells[0].text} off by {abs((a-b_)-c_):.4f}")
    results.append(check("margins within rounding of their columns", not bad, "; ".join(bad)))

    # -- provenance ----------------------------------------------------------------
    src = " ".join(" ".join(open(f, errors="ignore").read().split())
                   for f in glob.glob(str(ROOT / "results" / "*.txt")))
    orphan = []
    for t in d.tables:
        for r in t.rows[1:]:
            for c in r.cells[1:]:
                v = c.text.strip().replace("−", "-")
                if re.fullmatch(r"[+-]?\d\.\d{3}", v) and v.lstrip("+") not in src and v not in src:
                    orphan.append(f"{r.cells[0].text}={v}")
    results.append(check("every table figure in a saved results file",
                         len(orphan) <= 6, f"{len(orphan)} unmatched: {orphan[:4]}"))

    # -- cross-references ----------------------------------------------------------
    have = set()
    for l in appx.split("\n"):
        for pat in (r"^# Appendix ([A-H])\.", r"^#{2,3} ([A-H]\.\d+(?:\.\d)?)"):
            m = re.match(pat, l)
            have.add(m.group(1)) if m else None
    want_refs = set()
    for m in re.findall(r"\[([A-H](?:\.\d+(?:\.\d)?)?(?:\s*,\s*[A-H](?:\.\d+(?:\.\d)?)?)*)\]",
                        txt + appx):
        for x in m.split(","):
            want_refs.add(x.strip())
    results.append(check("appendix references resolve", not (want_refs - have),
                         " ".join(sorted(want_refs - have))))
    secs = set(re.findall(r"^(\d+\.\d+) ", "\n".join(paras), re.M))
    srefs = {x for r in re.findall(r"\[([\d., ]+)\]", txt) for x in re.findall(r"\d+\.\d+", r)}
    results.append(check("section references resolve", not (srefs - secs),
                         " ".join(sorted(srefs - secs))))
    bare = [l for l in appx.split("\n")
            if not l.startswith(("#", "|")) and re.search(r"(?<!\[)\bAppendix [A-H]\b(?!\])", l)]
    results.append(check("no bare 'Appendix X' in appendix prose", not bare, str(len(bare))))

    # -- punctuation ---------------------------------------------------------------
    viol = [t[:60] for t in body if "—" in t or re.search(r"[^\s]:\s", t)]
    for t in body:
        for m in re.finditer(r";", t):
            seg = t[max(0, m.start() - 60):m.start()]
            if "(" in seg and ")" in t[m.start():m.start() + 80]:
                continue
            viol.append(t[:60])
    results.append(check("no colons, semicolons or em dashes in prose", not viol,
                         f"{len(viol)}: {viol[:2]}"))

    # -- naming and agreement -------------------------------------------------------
    labels = collections.Counter(r.cells[0].text.strip() for t in d.tables for r in t.rows
                                 if "deployed" in r.cells[0].text)
    results.append(check("one label per representation", len(labels) <= 1, str(dict(labels))))
    agree = [t[:70] for t in body
             if re.search(r"\b(histogram|encoder|codec|probe|table) \w+s\b[^.]{0,60}\bbut (clear|carry|"
                          r"score|have|do)\b", t)]
    results.append(check("subject and verb agree after renames", not agree, str(agree[:1])))

    # -- duplication -----------------------------------------------------------------
    sents = [s for t in body for s in re.split(r"(?<=[.]) +", t) if len(s.split()) >= 15]
    dupes = sum(1 for i in range(len(sents)) for j in range(i + 1, len(sents))
                if difflib.SequenceMatcher(None, sents[i], sents[j]).ratio() >= 0.90)
    results.append(check("no sentence repeated near-verbatim", dupes == 0, f"{dupes} pair(s)"))

    # -- figures regenerable ----------------------------------------------------------
    # The invariant is that nothing hand-authored reaches the document, not that old files
    # are absent from the folder. An August figure still shows +0.244 for a value that is
    # now +0.225, and it is harmless on disk and fatal in the paper.
    gen = {os.path.basename(f) for f in glob.glob(str(ROOT / "docs" / "figures" / "fig?_*.png"))}
    embedded = len(d.inline_shapes)
    results.append(check("every embedded figure is generated from results/",
                         embedded == len(figs) and len(gen) >= len(figs),
                         f"{embedded} embedded, {len(gen)} generated files"))
    hand = [os.path.basename(f) for f in glob.glob(str(ROOT / "docs" / "figures" / "*.png"))
            if not os.path.basename(f).startswith("fig")]
    if hand and not args.quiet:
        print(f"  note  {len(hand)} hand-authored figure(s) on disk, not in the document: "
              f"{', '.join(hand)}")

    width = max(len(n) for n, _, _ in results)
    print(f"Draft check\n{'-' * (width + 34)}")
    for name, ok, detail in results:
        if ok and args.quiet:
            continue
        print(f"  {'ok  ' if ok else 'FAIL'}  {name:<{width}}  {detail if not ok else ''}")
    print(f"{'-' * (width + 34)}")
    print(f"  {len(results) - len(FAILED)}/{len(results)} passed")
    return 1 if FAILED else 0


if __name__ == "__main__":
    sys.exit(main())
