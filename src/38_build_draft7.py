#!/usr/bin/env python3
"""Build Draft 7 from Draft 6 by applying the citation audit fixes.

Replacements are applied at paragraph level, preserving the formatting of the first
affected run. Every change is reported so nothing happens silently.
"""
import os, shutil, sys
import docx

D = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/UCL Compling/Dissertation/Aug 6")
SRC = os.path.join(D, "Dissert Draft 6.docx")
DST = os.path.join(D, "Dissert Draft 7.docx")

REPLACEMENTS = [
    # --- 1. Qian suffix unified. Same paper cited two ways. ---
    ("Qian, Figueroa and Skantze (2025)", "Qian, Figueroa and Skantze (2025b)",
     "unify Qian suffix"),

    # --- 2. Gichamba downgraded from corroboration to characterisation. ---
    ("and attributing Mimi's performance at 12.5 Hz to its transformer bottleneck.",
     "and describing Mimi as engineered for low frame rate tokenisation through a "
     "transformer bottleneck and split-RVQ design. That description characterises the "
     "architecture rather than reporting a finding of their ablation, so it is consistent "
     "with the account here rather than independent evidence for it.",
     "downgrade Gichamba attribution"),

    # --- 3. Lotfian and Busso 2019 is not the paper on file. Busso et al. is. ---
    ("naturalistic podcast-derived corpora were built (Lotfian and Busso, 2019).",
     "naturalistic podcast-derived corpora were built (Busso et al., 2025).",
     "Lotfian and Busso -> Busso et al., first instance"),
    ("naturalistic podcast-derived corpora were built in response to exactly this "
     "(Lotfian and Busso, 2019)",
     "naturalistic podcast-derived corpora were built in response to exactly this "
     "(Busso et al., 2025)",
     "Lotfian and Busso -> Busso et al., second instance"),

    # --- 4. ABX: add the source that is actually held and verified. ---
    ("so it degrades gracefully at small samples (Schatz et al., 2013).",
     "so it degrades gracefully at small samples (Schatz et al., 2013), and it is used this "
     "way inside multi-level evaluation suites for spoken language modelling (Dunbar et al., "
     "2021).",
     "add Dunbar et al. 2021 alongside Schatz"),

    # --- 5. Near-verbatim repetition between 5.4 and 6.2. ---
    ("The architectural account is the best supported explanation available and it is not "
     "the only one. Gichamba and Busogi (2026) show an apparent architectural limit in DAC "
     "resolving into a training misconfiguration once sequence length was matched.",
     "The same caveat bounds the architectural account stated in [5.4]. Frozen public "
     "checkpoints cannot separate what an architecture can represent from what a particular "
     "training run taught it to represent, and the concrete reason for caution is that an "
     "apparent architectural limit in DAC has already been shown to resolve into a training "
     "misconfiguration (Gichamba and Busogi, 2026).",
     "de-duplicate 5.4 / 6.2"),
]


def replace_in_para(p, old, new):
    """Replace in a paragraph, keeping the first affected run's formatting."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return 0
    n = 0
    while old in full:
        start = full.index(old)
        end = start + len(old)
        pos, first, spans = 0, None, []
        for i, r in enumerate(p.runs):
            rs, re_ = pos, pos + len(r.text)
            if re_ > start and rs < end:
                if first is None:
                    first = i
                spans.append((i, rs, re_))
            pos = re_
        if first is None:
            break
        head = p.runs[first].text[: max(0, start - [s for i, s, e in spans if i == first][0])]
        tail_i, tail_s, tail_e = spans[-1]
        tail = p.runs[tail_i].text[end - tail_s:]
        p.runs[first].text = head + new + (tail if tail_i == first else "")
        for i, _, _ in spans[1:]:
            p.runs[i].text = "" if i != tail_i else tail
        full = "".join(r.text for r in p.runs)
        n += 1
    return n


def main():
    if not os.path.exists(SRC):
        sys.exit(f"missing {SRC}")
    d = docx.Document(SRC)
    paras = list(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                paras.extend(c.paragraphs)

    print(f"Draft 7 from Draft 6\n{'-'*70}")
    total = 0
    for old, new, label in REPLACEMENTS:
        hits = sum(replace_in_para(p, old, new) for p in paras)
        total += hits
        status = "ok " if hits else "** NOT FOUND **"
        print(f"  {status} {label:44} x{hits}")
    d.save(DST)
    print(f"{'-'*70}\n  {total} replacements, saved to {os.path.basename(DST)}")


if __name__ == "__main__":
    main()
