#!/usr/bin/env python3
"""Rebuild the Draft 7 tables that were incomplete or misordered."""
import os, copy
import docx

P = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/UCL Compling/Dissertation/Aug 6/Dissert Draft 7.docx")

T0_ADD = [  # Table 2.1, the two variable-rate tokenisers
 ("Sylber (Cho et al., 2025)", "16 kHz", "~4, variable", "", "", "syllabic, signal-derived units", "conv + transformer, HuBERT-based"),
 ("DyCAST (Della Libera et al., 2026)", "16 kHz", "6–24, variable", "", "4^32", "character-aligned, variable rate", "conv + transformer, WavLM-based"),
]

T1_ADD = [  # Table 2.2, prior work
 ("Pasad et al. (2021)", "read speech", "no", "no", "no", "layer-wise, phonetic and word"),
 ("Shi et al. (2026)", "read speech", "no", "yes", "partial, by codebook", "semantic and phonetic"),
 ("Sanders et al. (2025)", "acted and read", "no", "yes", "no", "paralinguistic and prosodic"),
 ("Pang et al. (2026)", "synthesised", "yes, system level", "n/a", "no", "behavioural"),
]

T2 = [  # Table 4.1, all fourteen, descending
 ("WavLM L20","0.557","0.009"),("Whisper encoder L9","0.548","0.013"),
 ("Text, target word only","0.493","0.004"),("Mimi, before quantisation","0.468","0.014"),
 ("Sylber","0.446","0.011"),("Mimi, after quantisation","0.441","0.014"),
 ("eGeMAPS, 88 functionals","0.420","0.012"),("DAC, before quantisation","0.404","0.015"),
 ("DyCAST, after quantisation","0.403","0.010"),("DyCAST, before quantisation","0.401","0.009"),
 ("EnCodec, after quantisation","0.400","0.010"),("EnCodec, before quantisation","0.396","0.012"),
 ("Text, with discourse context","0.394","0.013"),("DAC, after quantisation","0.381","0.013"),
 ("Mimi, deployed tokens","0.371","0.009"),
]

T3 = [  # nulls and margins, descending by margin
 ("WavLM L20","0.557","0.330","+0.226"),("Whisper encoder L9","0.548","0.333","+0.215"),
 ("Text, target word only","0.493","0.313","+0.179"),("Mimi, before quantisation","0.468","0.329","+0.139"),
 ("Sylber","0.446","0.328","+0.118"),("Mimi, after quantisation","0.441","0.337","+0.104"),
 ("eGeMAPS, 88 functionals","0.420","0.324","+0.096"),("DyCAST, after quantisation","0.403","0.317","+0.086"),
 ("DyCAST, before quantisation","0.401","0.320","+0.081"),("EnCodec, after quantisation","0.400","0.328","+0.072"),
 ("EnCodec, before quantisation","0.396","0.326","+0.070"),("DAC, before quantisation","0.404","0.336","+0.068"),
 ("Text, with discourse context","0.394","0.333","+0.061"),("Mimi, deployed tokens","0.371","0.311","+0.060"),
 ("DAC, after quantisation","0.381","0.331","+0.050"),
]

T6 = [  # cue retention, EnCodec inserted between Mimi and DAC
 ("WavLM, ceiling, absolute R²","0.283","0.607","0.225","0.557","0.345"),
 ("Mimi, before quantisation","153%","147%","158%","106%","170%"),
 ("Mimi, after quantisation","139%","143%","139%","97%","162%"),
 ("Mimi, deployed histogram","98%","106%","109%","78%","121%"),
 ("EnCodec, before quantisation","129%","145%","140%","72%","164%"),
 ("EnCodec, after quantisation","130%","144%","134%","71%","162%"),
 ("DAC, before quantisation","122%","130%","123%","68%","161%"),
 ("DAC, after quantisation","117%","127%","113%","63%","157%"),
]

T7 = [  # order effect, descending, all eleven
 ("WavLM L20","+0.113","0.013","44.6"),("Mimi, before quantisation","+0.080","0.018","21.9"),
 ("Whisper encoder L9","+0.070","0.015","23.6"),("EnCodec, after quantisation","+0.063","0.021","15.0"),
 ("Mimi, after quantisation","+0.048","0.019","13.0"),("Sylber","+0.042","0.016","13.1"),
 ("EnCodec, before quantisation","+0.033","0.017","9.8"),("DyCAST, before quantisation","+0.025","0.014","8.7"),
 ("DyCAST, after quantisation","+0.022","0.017","6.6"),("DAC, before quantisation","−0.007","0.019","−1.9"),
 ("DAC, after quantisation","−0.018","0.018","−4.9"),
]

T9 = [  # arousal control, eGeMAPS added
 ("WavLM L20","0.557","0.517","0.519"),("Whisper encoder L9","0.548","0.542","0.518"),
 ("Mimi, before quantisation","0.468","0.421","0.446"),("eGeMAPS, 88 functionals","0.420","0.389","0.409"),
 ("Mimi, deployed tokens","0.371","0.363","0.361"),
]


def rebuild(t, rows):
    """Clear all body rows and write `rows`, keeping the header."""
    while len(t.rows) > 1:
        t._tbl.remove(t.rows[-1]._tr)
    for vals in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, vals):
            c.paragraphs[0].add_run(v)


def append(t, rows):
    for vals in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, vals):
            c.paragraphs[0].add_run(v)


def main():
    d = docx.Document(P)
    print("Draft 7 tables\n" + "-"*58)
    append(d.tables[0], T0_ADD);  print(f"  Table 2.1  + Sylber, DyCAST         -> {len(d.tables[0].rows)-1} rows")
    append(d.tables[1], T1_ADD)
    # keep "This study" last
    t1 = d.tables[1]
    for i, r in enumerate(t1.rows):
        if r.cells[0].text.strip().startswith("This study"):
            t1._tbl.append(r._tr); break
    print(f"  Table 2.2  + Pasad, Shi, Sanders, Pang -> {len(t1.rows)-1} rows")
    rebuild(d.tables[2], T2);     print(f"  Table 4.1  + Sylber, DyCAST         -> {len(d.tables[2].rows)-1} rows")
    rebuild(d.tables[3], T3);     print(f"  nulls      + Sylber, DyCAST         -> {len(d.tables[3].rows)-1} rows")
    rebuild(d.tables[6], T6);     print(f"  cue ret.   + EnCodec                -> {len(d.tables[6].rows)-1} rows")
    rebuild(d.tables[7], T7);     print(f"  order      reordered, + 3 rows      -> {len(d.tables[7].rows)-1} rows")
    rebuild(d.tables[9], T9);     print(f"  arousal    + eGeMAPS                -> {len(d.tables[9].rows)-1} rows")
    d.save(P)
    print("-"*58 + "\n  saved")


if __name__ == "__main__":
    main()
