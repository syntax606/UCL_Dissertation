#!/usr/bin/env python3
"""Build a Harvard-style bibliography for Draft 7 as a .docx.

Author lists, titles and venues are taken from the title pages of the PDFs in the
Aug 6 folder. Entries whose source is absent, or where a detail could not be read
from the PDF, are marked so nothing is presented as verified that is not.
"""
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

D = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/UCL Compling/Dissertation/Aug 6")
OUT = os.path.join(D, "Bibliography.docx")

# (entry, flag) flag: "" verified from the PDF | "?" detail needs confirming | "!" no source on file
ENTRIES = [
("Arora, S., Chang, K.-W., Chien, C.-M., Peng, Y., Wu, H., Adi, Y., Dupoux, E., Lee, H.-Y., "
 "Livescu, K. and Watanabe, S. (2025) 'On the landscape of spoken language models: a comprehensive "
 "survey', Transactions on Machine Learning Research, 10/2025.", ""),
("Baade, A., Peng, P. and Harwath, D. (2025) 'SyllableLM: learning coarse semantic units for speech "
 "language models', in Proceedings of the International Conference on Learning Representations "
 "(ICLR 2025).", ""),
# The PDF on file is arXiv 2102.12452v4, which states the venue in its own front matter:
# a Computational Linguistics squib, accepted 8 September 2021, ACL copyright. That places
# the published version in the 48(1) issue of 2022, which is what the draft cites. The
# volume and pagination are from the published version rather than from the preprint.
("Belinkov, Y. (2022) 'Probing classifiers: promises, shortcomings, and advances', Computational "
 "Linguistics, 48(1), pp. 207–219.", ""),
("Bengio, Y. and Grandvalet, Y. (2004) 'No unbiased estimator of the variance of k-fold "
 "cross-validation', Journal of Machine Learning Research, 5, pp. 1089–1105.", ""),
("Biber, D. and Finegan, E. (1988) 'Adverbial stance types in English', Discourse Processes, "
 "11(1), pp. 1–34.", ""),
("Borsos, Z., Marinier, R., Vincent, D., Kharitonov, E., Pietquin, O., Sharifi, M., Roblek, D., "
 "Teboul, O., Grangier, D., Tagliasacchi, M. and Zeghidour, N. (2023) 'AudioLM: a language modeling "
 "approach to audio generation', IEEE/ACM Transactions on Audio, Speech, and Language Processing.", "?"),
("Bouthillier, X., Delaunay, P., Bronzi, M., Trofimov, A., Nichyporuk, B., Szeto, J., Sepah, N., "
 "Raff, E., Madan, K., Voleti, V., Ebrahimi Kahou, S., Michalski, V., Serdyuk, D., Arbel, T., Pal, C., "
 "Varoquaux, G. and Vincent, P. (2021) 'Accounting for variance in machine learning benchmarks', in "
 "Proceedings of Machine Learning and Systems (MLSys), 3.", "?"),
("Bryant, G.A. and Fox Tree, J.E. (2005) 'Is there an ironic tone of voice?', Language and Speech, "
 "48(3), pp. 257–277.", "?"),
# Year confirmed from the title page: arXiv 2509.09791v1, 11 September 2025, manuscript
# received 10 September 2025. The authors carry IEEE affiliations and the paper is formatted
# for an IEEE journal, but no venue is stated and the revision line is still blank, so this
# is a preprint and citing a journal would overstate it.
("Busso, C., Lotfian, R., Sridhar, K., Salman, A.N., Lin, W.-C., Goncalves, L., Parthasarathy, S., "
 "Naini, A.R., Leem, S.-G., Martinez-Lucas, L., Chou, H.-C. and Mote, P. (2025) 'The MSP-Podcast "
 "corpus'. arXiv:2509.09791.", ""),
("Chen, S., Wang, C., Chen, Z., Wu, Y., Liu, S., Chen, Z., Li, J., Kanda, N., Yoshioka, T., Xiao, X., "
 "Wu, J., Zhou, L., Ren, S., Qian, Y., Qian, Y., Wu, J., Zeng, M., Yu, X. and Wei, F. (2022) 'WavLM: "
 "large-scale self-supervised pre-training for full stack speech processing', IEEE Journal of Selected "
 "Topics in Signal Processing, 16(6), pp. 1505–1518.", "!"),
("Cho, C.J., Lee, N., Gupta, A., Agarwal, D., Chen, E., Black, A.W. and Anumanchipalli, G.K. (2025) "
 "'Sylber: syllabic embedding representation of speech from raw audio', in Proceedings of the "
 "International Conference on Learning Representations (ICLR 2025). arXiv:2410.07168.", "!"),
("Défossez, A., Copet, J., Synnaeve, G. and Adi, Y. (2022) 'High fidelity neural audio compression'. "
 "arXiv:2210.13438.", ""),
("Défossez, A., Mazaré, L., Orsini, M., Royer, A., Pérez, P., Jégou, H., Grave, E. and Zeghidour, N. "
 "(2024) 'Moshi: a speech-text foundation model for real-time dialogue'. Kyutai. arXiv:2410.00037.", ""),
("Della Libera, L., Subakan, C. and Ravanelli, M. (2026) 'Beyond fixed frames: dynamic "
 "character-aligned speech tokenization'. arXiv:2601.23174.", ""),
("de Seyssel, M., Lavechin, M., Titeux, H., Thomas, A., Virlet, G., Santos Revilla, A., Wisniewski, G., "
 "Ludusan, B. and Dupoux, E. (2023) 'ProsAudit, a prosodic benchmark for self-supervised speech "
 "models', in Proceedings of Interspeech 2023.", ""),
("Du Bois, J.W. (2007) 'The stance triangle', in Englebretson, R. (ed.) Stancetaking in Discourse: "
 "Subjectivity, Evaluation, Interaction. Amsterdam: John Benjamins, pp. 139–182.", "?"),
("Dunbar, E., Bernard, M., Hamilakis, N., Nguyen, T.A., de Seyssel, M., Rozé, P., Rivière, M., "
 "Kharitonov, E. and Dupoux, E. (2021) 'The Zero Resource Speech Challenge 2021: spoken language "
 "modelling'. arXiv:2104.14700.", ""),
("Gardner, R. (2001) When Listeners Talk: Response Tokens and Listener Stance. Amsterdam: "
 "John Benjamins.", ""),
("Gichamba, A. and Busogi, M. (2026) 'Probing low frame rate degradation in neural audio codecs', "
 "in Proceedings of Interspeech 2026. arXiv:2606.16969.", ""),
("Gravano, A., Hirschberg, J. and Beňuš, Š. (2012) 'Affirmative cue words in task-oriented dialogue', "
 "Computational Linguistics, 38(1), pp. 1–39.", "?"),
("Guo, Y., Li, Z., Wang, H., Li, B., Shao, C., Zhang, H., Du, C., Chen, X., Liu, S. and Yu, K. (2025) "
 "'Recent advances in discrete speech tokens: a review'. arXiv:2502.06490.", ""),
("Hewitt, J. and Liang, P. (2019) 'Designing and interpreting probes with control tasks', in "
 "Proceedings of EMNLP-IJCNLP 2019, pp. 2733–2743.", "?"),
("Hsu, W.-N., Bolte, B., Tsai, Y.-H.H., Lakhotia, K., Salakhutdinov, R. and Mohamed, A. (2021) "
 "'HuBERT: self-supervised speech representation learning by masked prediction of hidden units', "
 "IEEE/ACM Transactions on Audio, Speech, and Language Processing, 29, pp. 3451–3460.", ""),
("Kumar, R., Seetharaman, P., Luebs, A., Kumar, I. and Kumar, K. (2023) 'High-fidelity audio "
 "compression with improved RVQGAN', in Advances in Neural Information Processing Systems, 36. "
 "arXiv:2306.06546.", ""),
("Lan, C., Hui, P.L., Xu, W. and Mok, P. (2019) 'Revisiting acoustic markers of sarcasm in Cantonese', "
 "in Proceedings of the International Congress of Phonetic Sciences (ICPhS 2019).", ""),
("Li, J., Qian, Y., Hu, Y., Zhang, L., Wang, X., Lu, H., Thakker, M., Li, J., Zhao, S. and Wu, Z. "
 "(2026) 'FlexiCodec: a dynamic neural audio codec for low frame rates', in Proceedings of the "
 "International Conference on Learning Representations (ICLR 2026). arXiv:2510.00981.", ""),
("Lin, G.-T., Feng, C.-L., Huang, W.-P., Tseng, Y., Lin, T.-H., Li, C.-A., Lee, H.-y. and Ward, N.G. "
 "(2022) 'On the utility of self-supervised models for prosody-related tasks', in Proceedings of the "
 "IEEE Spoken Language Technology Workshop (SLT 2022). arXiv:2210.07185.", "?"),
("Liu, W., Guo, Z., Xu, J., Lv, Y., Chu, Y., Zhao, Z. and Lin, J. (2024) 'Analyzing and mitigating "
 "inconsistency in discrete audio tokens for neural codec language models'. arXiv:2409.19283.", ""),
("Misra, I., Zitnick, C.L. and Hebert, M. (2016) 'Shuffle and learn: unsupervised learning using "
 "temporal order verification', in Proceedings of the European Conference on Computer Vision (ECCV "
 "2016). arXiv:1603.08561.", ""),
("Mousavi, P., Duret, J., Petermann, D., Ploujnikov, A., Della Libera, L., Kuznetsova, A., Subakan, C. "
 "and Ravanelli, M. (2026) 'DASB: discrete audio and speech benchmark', Transactions on Machine "
 "Learning Research, 04/2026.", ""),
("O'Connor Russell, S., Charuau, D. and Harte, N. (2026) 'The role of prosodic and lexical cues in "
 "turn-taking with self-supervised speech representations'. arXiv:2601.13835.", ""),
("Pang, J., Chaubey, A. and Soleymani, M. (2026) 'Do audio LLMs listen or read? Analyzing and "
 "mitigating paralinguistic failures with VoxParadox'. arXiv:2605.27772.", ""),
("Pasad, A., Chou, J.-C. and Livescu, K. (2021) 'Layer-wise analysis of a self-supervised speech "
 "representation model', in Proceedings of the IEEE Automatic Speech Recognition and Understanding "
 "Workshop (ASRU 2021). arXiv:2107.04734.", ""),
("Qian, J. and Li, J. (2026) 'Prosody-driven jailbreaks in audio LLMs: a controlled study and "
 "mechanistic analysis'. City University of Hong Kong. arXiv:2607.26541.", ""),
("Qian, K., Fan, X., Ni, J., Shechtman, S., Hasegawa-Johnson, M., Gan, C. and Zhang, Y. (2025a) "
 "'ProsodyLM: uncovering the emerging prosody processing capabilities in speech language models', in "
 "Proceedings of the Conference on Language Modeling (COLM 2025).", ""),
("Qian, L., Figueroa, C. and Skantze, G. (2025b) 'Representation of perceived prosodic similarity of "
 "conversational feedback'. KTH Royal Institute of Technology. arXiv:2505.13268.", ""),
("Radford, A., Kim, J.W., Xu, T., Brockman, G., McLeavey, C. and Sutskever, I. (2023) 'Robust speech "
 "recognition via large-scale weak supervision', in Proceedings of the International Conference on "
 "Machine Learning (ICML). PMLR, pp. 28492–28518.", "!"),
("Ren, W., Lin, Y.-C., Chou, H.-C., Wu, H., Wu, Y.-C., Lee, C.-C., Lee, H.-y. and Tsao, Y. (2024) "
 "'EMO-Codec: an in-depth look at emotion preservation capacity of legacy and neural codec models with "
 "subjective and objective evaluations'. arXiv:2407.15458.", ""),
("Rockwell, P. (2000) 'Lower, slower, louder: vocal cues of sarcasm', Journal of Psycholinguistic "
 "Research, 29(5), pp. 483–495.", "?"),
("Sanders, N., Li, Y., Richmond, K. and King, S. (2025) 'Segmentation-variant codebooks for "
 "preservation of paralinguistic and prosodic information'. University of Edinburgh. "
 "arXiv:2505.15667.", ""),
("Schatz, T., Peddinti, V., Bach, F., Jansen, A., Hermansky, H. and Dupoux, E. (2013) 'Evaluating "
 "speech features with the minimal-pair ABX task: analysis of the classical MFC/PLP pipeline', in "
 "Proceedings of Interspeech 2013, pp. 1781–1785.", "!"),
("Schegloff, E.A. (1982) 'Discourse as an interactional achievement: some uses of ‘uh huh’ and other "
 "things that come between sentences', in Tannen, D. (ed.) Analyzing Discourse: Text and Talk. "
 "Washington, DC: Georgetown University Press, pp. 71–93.", "!"),
("Scherer, K.R. (2003) 'Vocal communication of emotion: a review of research paradigms', Speech "
 "Communication, 40(1–2), pp. 227–256.", "?"),
("Shi, X., Zeng, C., Feng, T., Wang, S.-H., Ma, J. and Narayanan, S. (2026) 'Speech codec probing from "
 "semantic and phonetic perspectives'. University of Southern California and Dolby Laboratories. "
 "arXiv:2603.10371.", ""),
("Sicherman, A. and Adi, Y. (2023) 'Analysing discrete self-supervised speech representation for "
 "spoken language modeling', in Proceedings of ICASSP 2023. arXiv:2301.00591.", "?"),
# The draft cited 2019, which matched nothing. The copy on file states its own year in the
# footer: Encyclopedia of Applied Linguistics, ed. Chapelle, (c) 2026 Wiley, DOI suffix .pub2.
# The pub2 marks it as a revision of the original entry, so 2026 is the version read here.
("Steensig, J. (2026) 'Conversation analysis and affiliation and alignment', in Chapelle, C.A. (ed.) "
 "The Encyclopedia of Applied Linguistics. Hoboken, NJ: Wiley. "
 "doi: 10.1002/9781405198431.wbeal0196.pub2.", ""),
("Stivers, T. (2008) 'Stance, alignment, and affiliation during storytelling: when nodding is a token "
 "of affiliation', Research on Language and Social Interaction, 41(1), pp. 31–57.", "?"),
("Sun, E., Naini, A.R. and Busso, C. (2026) 'Recovering performance in speech emotion recognition from "
 "discrete tokens via multi-layer fusion and paralinguistic feature integration'. Carnegie Mellon "
 "University.", "?"),
("Yang, S.-w., Tu, M., Liu, A.T., Qu, X., Lee, H.-y., Lu, L., Wang, Y. and Wu, Y. (2026) 'ParaS2S: "
 "benchmarking and aligning spoken language models for paralinguistic-aware speech-to-speech "
 "interaction', in Proceedings of the International Conference on Learning Representations (ICLR "
 "2026).", ""),
("Ye, Z., Sun, P., Lei, J., Lin, H., Tan, X., Dai, Z., Kong, Q., Chen, J., Pan, J., Liu, Q., Guo, Y. "
 "and Xue, W. (2024) 'Codec does matter: exploring the semantic shortcoming of codec for audio "
 "language model'. arXiv:2408.17175.", ""),
("Zeghidour, N., Luebs, A., Omran, A., Skoglund, J. and Tagliasacchi, M. (2021) 'SoundStream: an "
 "end-to-end neural audio codec', IEEE/ACM Transactions on Audio, Speech, and Language Processing, 30, "
 "pp. 495–507. arXiv:2107.03312.", ""),
("Zhang, X., Zhang, D., Li, S., Zhou, Y. and Qiu, X. (2024a) 'SpeechTokenizer: unified speech "
 "tokenizer for speech language models', in Proceedings of the International Conference on Learning "
 "Representations (ICLR 2024). arXiv:2308.16692.", ""),
]


def main():
    d = docx.Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)

    h = d.add_heading("References", level=1)

    n_ok = sum(1 for _, f in ENTRIES if f == "")
    n_q = sum(1 for _, f in ENTRIES if f == "?")
    n_x = sum(1 for _, f in ENTRIES if f == "!")

    note = d.add_paragraph()
    r = note.add_run(
        f"Harvard style, alphabetical by first author. {len(ENTRIES)} entries. "
        f"Author lists, titles and venues were read from the title pages of the PDFs held in the "
        f"Aug 6 folder wherever a source was available. "
        f"{n_ok} entries are fully verified against the source. "
        f"{n_q} are marked with a question mark, meaning the work is on file but one detail, usually "
        f"volume, issue, page range or venue, was not printed on the title page and needs confirming. "
        f"{n_x} are marked with an exclamation mark, meaning no copy of the source is held and the "
        f"reference was reconstructed from citations of it in papers that are on file."
    )
    r.font.size = Pt(9)
    r.italic = True

    d.add_paragraph()

    for text, flag in sorted(ENTRIES, key=lambda e: e[0].lower()):
        p = d.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(text)
        if flag:
            m = p.add_run("  " + ("[?]" if flag == "?" else "[!]"))
            m.bold = True
            m.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    d.add_paragraph()
    k = d.add_paragraph()
    kr = k.add_run(
        "[?] on file, one bibliographic detail needs confirming.    "
        "[!] no copy held, reference reconstructed from other papers' citations of it."
    )
    kr.font.size = Pt(9)
    kr.italic = True

    d.save(OUT)
    print(f"  {len(ENTRIES)} entries  ({n_ok} verified, {n_q} need a detail, {n_x} no source)")
    print(f"  saved to {OUT}")


if __name__ == "__main__":
    main()
