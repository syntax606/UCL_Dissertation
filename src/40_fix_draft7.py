#!/usr/bin/env python3
"""Apply the Draft 7 review fixes in place. Every change is reported."""
import os, copy
import docx

D = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/UCL Compling/Dissertation/Aug 6")
P = os.path.join(D, "Dissert Draft 7.docx")

# full null table, all twelve representations, measured under the scheme in 3.7
NULLS = [
    ("WavLM L20", "0.557", "0.330", "+0.226"),
    ("Whisper encoder L9", "0.548", "0.333", "+0.215"),
    ("Text, target word only", "0.493", "0.313", "+0.179"),
    ("Mimi, before quantisation", "0.468", "0.329", "+0.139"),
    ("Mimi, after quantisation", "0.441", "0.337", "+0.104"),
    ("eGeMAPS, 88 functionals", "0.420", "0.324", "+0.096"),
    ("EnCodec, after quantisation", "0.400", "0.328", "+0.072"),
    ("EnCodec, before quantisation", "0.396", "0.326", "+0.070"),
    ("DAC, before quantisation", "0.404", "0.336", "+0.068"),
    ("Text, with discourse context", "0.394", "0.333", "+0.061"),
    ("Mimi, deployed tokens", "0.371", "0.311", "+0.060"),
    ("DAC, after quantisation", "0.381", "0.331", "+0.050"),
]

TEXT_FIXES = [
    # --- repeated clause inside 6.2 ---
    ("Frozen public checkpoints cannot separate what an architecture can represent from what a "
     "particular training run taught it to represent, and settling that needs codecs trained under "
     "matched conditions differing only in the temporal mechanism.",
     "Settling it needs codecs trained under matched conditions differing only in the temporal "
     "mechanism.",
     "6.2 repeated clause"),

    # --- constraint count: four paragraphs follow, and the identity control is its own point ---
    ("Five constraints bound the positive findings.",
     "Four constraints bound the positive findings, and a fifth concerns the identity control.",
     "6.2 constraint count"),

    # --- null range corrected and coverage claim made true ---
    ("The null is not constant across configurations and ranges from 0.311 to 0.333, falling lowest "
     "for the 16,384-dimensional histogram, so margins are given against each configuration's own "
     "null rather than against a single assumed value [B.2].",
     "The null is not constant across configurations, ranging from 0.311 to 0.337 and falling lowest "
     "for the 16,384-dimensional histogram, so every margin below is given against that "
     "configuration's own null rather than against a single assumed value [B.2]. One consequence is "
     "visible in the table. Mimi's deployed tokens score lower than DAC after quantisation, at 0.371 "
     "against 0.381, but clear a lower null and so carry the larger margin, at +0.060 against +0.050.",
     "4.2 null range and coverage"),
    ("ranging from 0.311 for the 16,384-dimensional histogram to 0.333 for the continuous encoders [4.2].",
     "ranging from 0.311 for the 16,384-dimensional histogram to 0.337 for Mimi's post-quantisation "
     "vectors [4.2].",
     "3.7 null range"),

    # --- 6.1 flattens the two-round hypothesis structure ---
    ("The four research questions are answered, and the hypotheses stated in [1.3] resolve as follows: "
     "H1, H2, H3, H5 and H9 are supported, while H4 and H6 to H8 are not.",
     "The four research questions are answered. Of the five hypotheses formed before the analyses, H1, "
     "H2, H3 and H5 are supported and H4 is not. Of the four formed afterwards, once the falsification "
     "of H4 had redirected the study, H9 is supported and H6 to H8 are not. That H9 was formed after an "
     "anomalous result and then tested on a codec not previously included is stated here as it is in "
     "[1.3], since a hypothesis confirmed on new data is not the same kind of claim as one registered "
     "in advance.",
     "6.1 two-round hypothesis structure"),

    # --- 6.1 broken sentence ---
    ("Every figure above is a mean over 25 partitions for that reason, and the practice costs minutes. "
     "so differences under roughly 0.03 are not read as orderings.",
     "Every figure above is a mean over 25 partitions for that reason, and the practice costs minutes, "
     "so differences under roughly 0.03 are not read as orderings.",
     "6.1 broken sentence"),

    # --- 4.4 digressions moved out of the main line ---
    ("One figure in that table is layer-sensitive and should be read with the sweep. Whisper is reported "
     "at L9, the layer fixed in advance [3.7], where the order effect is +0.070. At L12 it reaches "
     "+0.117, which would place it above WavLM [G.6]. The stance decoding barely moves between those "
     "layers, at 0.548 against 0.551, so the sensitivity is specific to the order measurement. Reporting "
     "L9 follows the pre-fixed rule rather than the larger figure, and the codec ladder is unaffected "
     "either way.",
     "Whisper's figure is layer-sensitive in a way the others are not, and the sweep in [G.6] should be "
     "read alongside it.",
     "4.4 move Whisper digression"),
]

# 2.4 gains back the words that buys
EXPAND_ANCHOR = ("Response tokens. Yeah, okay, right and sure are response tokens. Schegloff")
EXPAND_INSERT = (
    "The construct these labels name is not a convenience of this study. Stance in Du Bois's sense is a "
    "single act with three faces, evaluating an object, positioning the speaker and calibrating "
    "alignment with an interlocutor, and it is the third that a response token performs almost to the "
    "exclusion of the other two. Biber and Finegan's account of stance marking is lexical and "
    "grammatical, which is precisely what the design removes by holding the word constant, so what "
    "remains when those markers are gone is the interactional layer that conversation analysis "
    "describes as affiliation and disaffiliation. Stivers shows that layer being negotiated turn by "
    "turn during storytelling, where a nod can affiliate or withhold affiliation without any lexical "
    "content at all, and Steensig separates alignment, which concerns the structural progress of the "
    "activity, from affiliation, which concerns endorsement of the stance being displayed. The "
    "three-way axis annotated here maps onto affiliation rather than alignment, which is why a "
    "structurally cooperative backchannel can still be adversarial, and why the neutral category is "
    "not an absence of stance but a distinct interactional move."
)


def replace_in_para(p, old, new):
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return 0
    start = full.index(old); end = start + len(old)
    pos, first, spans = 0, None, []
    for i, r in enumerate(p.runs):
        rs, re_ = pos, pos + len(r.text)
        if re_ > start and rs < end:
            if first is None: first = i
            spans.append((i, rs, re_))
        pos = re_
    if first is None: return 0
    fs = [s for i, s, e in spans if i == first][0]
    head = p.runs[first].text[: start - fs]
    ti, ts, te = spans[-1]
    tail = p.runs[ti].text[end - ts:]
    p.runs[first].text = head + new + (tail if ti == first else "")
    for i, _, _ in spans[1:]:
        p.runs[i].text = "" if i != ti else tail
    return 1


def main():
    d = docx.Document(P)
    paras = list(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                paras.extend(c.paragraphs)

    print(f"Draft 7 fixes\n{'-'*66}")

    # 1. duplicate heading
    killed = 0
    for i in range(len(d.paragraphs) - 1, 0, -1):
        a, b = d.paragraphs[i].text.strip(), d.paragraphs[i-1].text.strip()
        if a and a == b and a.startswith("6.1"):
            d.paragraphs[i]._element.getparent().remove(d.paragraphs[i]._element)
            killed += 1
    print(f"  {'ok ' if killed else '** none **'} duplicate '6.1 Conclusions' heading   x{killed}")

    # 2. text fixes
    for old, new, label in TEXT_FIXES:
        hits = sum(replace_in_para(p, old, new) for p in paras)
        print(f"  {'ok ' if hits else '** NOT FOUND **'} {label:40} x{hits}")

    # 3. expand 2.4
    done = 0
    for p in d.paragraphs:
        if p.text.strip().startswith(EXPAND_ANCHOR[:40]):
            new_p = copy.deepcopy(p._element)
            p._element.addprevious(new_p)
            np_ = docx.text.paragraph.Paragraph(new_p, p._parent)
            for r in np_.runs[1:]:
                r.text = ""
            np_.runs[0].text = EXPAND_INSERT
            done = 1
            break
    print(f"  {'ok ' if done else '** NOT FOUND **'} 2.4 expanded on the stance construct")

    # 4. Table 2.2 de Seyssel row
    fixed = 0
    for row in d.tables[1].rows:
        c = row.cells
        if "Seyssel" in c[0].text:
            vals = ["de Seyssel et al. (2023)", "read speech", "by construction", "no", "no",
                    "prosodic boundary and pause"]
            for cell, v in zip(c, vals):
                for pp in cell.paragraphs[1:]:
                    pp._element.getparent().remove(pp._element)
                para = cell.paragraphs[0]
                if para.runs:
                    para.runs[0].text = v
                    for r in para.runs[1:]: r.text = ""
                else:
                    para.add_run(v)
            fixed = 1
    print(f"  {'ok ' if fixed else '** NOT FOUND **'} Table 2.2 de Seyssel row completed")

    # 5. null table extended to all twelve
    t = d.tables[3]
    while len(t.rows) > 1:
        t._tbl.remove(t.rows[-1]._tr)
    for rep, f1, null, margin in NULLS:
        cells = t.add_row().cells
        for cell, v in zip(cells, (rep, f1, null, margin)):
            cell.paragraphs[0].add_run(v)
    print(f"  ok  null table extended from 5 to {len(t.rows)-1} representations")

    d.save(P)
    print(f"{'-'*66}\n  saved in place")


if __name__ == "__main__":
    main()
