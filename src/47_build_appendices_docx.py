#!/usr/bin/env python3
"""Build the appendices Word document from docs/drafts/07_appendices.md.

Direction of authority is the opposite of the chapters. For Chapters 1 to 6 the .docx is
authoritative and src/45_export_draft.py generates the markdown from it. The appendices
have no docx counterpart and have only ever existed as markdown, so here the markdown is
authoritative and this script generates the .docx from it. Edit the markdown, then re-run.

Styling is taken from the draft itself rather than declared, by opening Draft 7, keeping
its styles and table formatting, and clearing the body. Headings therefore match the
chapters at 26 pt, 19.5 pt and 16 pt bold, and body text uses the same p1 style.

Usage:  python3 src/47_build_appendices_docx.py
        python3 src/47_build_appendices_docx.py --out /some/other/path.docx
"""
import argparse
import copy
import os
import re
from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "drafts" / "07_appendices.md"
AUG6 = os.path.expanduser("~/Library/Mobile Documents/com~apple~CloudDocs/"
                          "UCL Compling/Dissertation/Aug 6")
TEMPLATE = os.path.join(AUG6, "Dissert Draft 7.docx")
OUT = os.path.join(AUG6, "Dissert Draft 7 Appendices.docx")

BODY_PT = 15.0     # p1 body text in the chapters
H1_PT = 26.0       # "Chapter 1: Introduction"
H2_PT = 19.5       # "1.1 The same word..."
H3_PT = 16.0       # no chapter equivalent, sits between H2 and a table caption
TABLE_PT = 13.0    # table cells and captions in the chapters
FONT = "UICTFontTextStyleBody"
MONO = "Courier New"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<![A-Za-z0-9])_[^_]+_(?![A-Za-z0-9]))")


def add_runs(p, text, *, size=BODY_PT, bold=False, italic=False, mono=False):
    """Write text into a paragraph, honouring **bold**, `code` and _italic_.

    Recursive, because the markdown nests. G.6 has a run-in heading of the form
    **`seg4` is the strongest readout**, and a single pass would strip the bold markers
    and leave the backticks sitting in the text.
    """
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            add_runs(p, tok[2:-2], size=size, bold=True, italic=italic, mono=mono)
        elif tok.startswith("`") and tok.endswith("`"):
            add_runs(p, tok[1:-1], size=size, bold=bold, italic=italic, mono=True)
        elif tok.startswith("_") and tok.endswith("_"):
            add_runs(p, tok[1:-1], size=size, bold=bold, italic=True, mono=mono)
        else:
            r = p.add_run(tok)
            r.bold, r.italic = bold, italic
            r.font.size = Pt(size)
            r.font.name = MONO if mono else FONT
    return p


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_rule(line):
    return bool(re.match(r"^\|[\s:|-]+\|?$", line.strip())) and "-" in line


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=OUT)
    ap.add_argument("--template", default=TEMPLATE)
    args = ap.parse_args()

    doc = docx.Document(args.template)

    # Keep one table's formatting before the body is cleared, so the appendix tables
    # look like the chapter tables rather than like the bare default.
    tbl_pr = copy.deepcopy(doc.tables[0]._tbl.tblPr) if doc.tables else None
    body = doc.element.body
    for child in list(body.iterchildren()):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)

    lines = SRC.read_text().split("\n")
    n_tables = n_head = 0
    first_appendix = True
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        # table: a header row followed by a |---| rule
        if line.startswith("|") and i + 1 < len(lines) and is_rule(lines[i + 1]):
            header = split_row(line)
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            if tbl_pr is not None:
                t._tbl.remove(t._tbl.tblPr)
                t._tbl.insert(0, copy.deepcopy(tbl_pr))
            for cell, txt in zip(t.rows[0].cells, header):
                add_runs(cell.paragraphs[0], txt, size=TABLE_PT, bold=True)
            for r in rows:
                cells = t.add_row().cells
                for cell, txt in zip(cells, r[: len(header)]):
                    add_runs(cell.paragraphs[0], txt, size=TABLE_PT)
            n_tables += 1
            continue

        m = re.match(r"^(#{1,3}) +(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            if level == 1 and text.startswith("Appendix"):
                if not first_appendix:
                    p.add_run().add_break(WD_BREAK.PAGE)
                first_appendix = False
            add_runs(p, text, size={1: H1_PT, 2: H2_PT, 3: H3_PT}[level], bold=True)
            n_head += 1
            i += 1
            continue

        # a table caption in the chapters is bold at table size
        if re.match(r"^(Table|Figure) [A-H0-9]", line.strip()):
            add_runs(doc.add_paragraph(), line.strip(), size=TABLE_PT, bold=True)
            i += 1
            continue

        # a wholly italic line, used for the note under the title
        if line.strip().startswith("*(") and line.strip().endswith(")*"):
            buf = [line.strip()]
            while not buf[-1].endswith(")*") and i + 1 < len(lines):
                i += 1
                buf.append(lines[i].strip())
            add_runs(doc.add_paragraph(), " ".join(buf)[1:-1], italic=True)
            i += 1
            continue

        # ordinary paragraph, joining its wrapped lines
        buf = [line.strip()]
        while (i + 1 < len(lines) and lines[i + 1].strip()
               and not lines[i + 1].startswith(("#", "|", "---"))):
            i += 1
            buf.append(lines[i].strip())
        p = doc.add_paragraph(style="p1")
        add_runs(p, " ".join(buf))
        i += 1

    doc.save(args.out)
    words = sum(len(p.text.split()) for p in docx.Document(args.out).paragraphs)
    print(f"built {os.path.basename(args.out)}\n{'-' * 58}")
    print(f"  {n_head} headings, {n_tables} tables, {words} words outside tables")
    print(f"  from {SRC.relative_to(ROOT)}, styled from {os.path.basename(args.template)}")
    print(f"  -> {args.out}")
    print("\n  the markdown is authoritative for the appendices; re-run this after editing it")


if __name__ == "__main__":
    main()
